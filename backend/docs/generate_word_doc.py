"""
Generate a structured Word document from the Super-Agent project documentation.

Usage:
    python generate_word_doc.py

Output:
    super_agent_documentation.docx  (in the same directory)
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# ── Ordered list of (title, filename) for the document sections ──────────────
SECTIONS = [
    ("System Analysis",                  "system_analysis.md"),
    ("Feature Suggestions",              "feature_suggestions.md"),
    ("Universal Feature Proposals",      "universal_features.md"),
    ("Unprecedented Feature: AESCE",     "unprecedented_feature.md"),
    ("AESCE Implementation Plan",        "implementation_plan.md"),
    ("AESCE Walkthrough",                "walkthrough.md"),
    ("Implementation Task Checklist",    "task.md"),
]

DOCS_DIR = Path(__file__).parent
OUTPUT_FILE = DOCS_DIR / "super_agent_documentation.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _add_toc(doc: Document) -> None:
    """Insert a Word-native Table of Contents field."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def _set_run_color(run, hex_color: str) -> None:
    """Set run font color using a hex string like '1F3864'."""
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _add_page_break(doc: Document) -> None:
    doc.add_page_break()


def _parse_inline(paragraph, text: str) -> None:
    """
    Parse inline markdown (bold **text**, inline code `text`) and add
    appropriately formatted runs to *paragraph*.
    """
    # Split on bold (**…**) and inline-code (`…`) markers
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)


def _add_markdown_section(doc: Document, markdown_text: str) -> None:
    """
    Convert a block of markdown text into Word paragraphs/headings.

    Supported elements:
    - # / ## / ### headings  → Heading 1/2/3
    - - [ ] / - [x] task lists → styled list items
    - Bullet lists (- item, * item)
    - Numbered lists (1. item)
    - Fenced code blocks (``` … ```)
    - Blank lines → paragraph breaks
    - Inline **bold** and `code`
    """
    in_code_block = False
    code_lines: list[str] = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        # ── Fenced code blocks ─────────────────────────────────────────────
        if line.startswith("```"):
            if in_code_block:
                # Close the block — emit collected code
                code_para = doc.add_paragraph("\n".join(code_lines), style="Normal")
                code_para.runs[0].font.name = "Courier New"
                code_para.runs[0].font.size = Pt(9)
                code_para.paragraph_format.left_indent = Inches(0.5)
                # Add a light background shading via XML
                pPr = code_para._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                pPr.append(shd)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        # Reset numbered-list counter on blank line
        if not line:
            doc.add_paragraph()
            continue

        # ── Headings ───────────────────────────────────────────────────────
        h3 = re.match(r"^#{3}\s+(.*)", line)
        h2 = re.match(r"^#{2}\s+(.*)", line)
        h1 = re.match(r"^#\s+(.*)", line)

        if h3:
            p = doc.add_heading(h3.group(1), level=3)
            continue
        if h2:
            p = doc.add_heading(h2.group(1), level=2)
            continue
        if h1:
            # Within a section the top-level heading becomes Heading 2
            p = doc.add_heading(h1.group(1), level=2)
            continue

        # ── Task-list items ────────────────────────────────────────────────
        task_done = re.match(r"^(\s*)-\s+\[x\]\s+(.*)", line, re.IGNORECASE)
        task_todo = re.match(r"^(\s*)-\s+\[ \]\s+(.*)", line, re.IGNORECASE)
        if task_done:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run("✔ ")
            run.bold = True
            _set_run_color(run, "2E7D32")
            _parse_inline(p, task_done.group(2))
            indent = len(task_done.group(1))
            if indent:
                p.paragraph_format.left_indent = Inches(indent * 0.25)
            continue
        if task_todo:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run("☐ ")
            _parse_inline(p, task_todo.group(2))
            indent = len(task_todo.group(1))
            if indent:
                p.paragraph_format.left_indent = Inches(indent * 0.25)
            continue

        # ── Bullet lists ───────────────────────────────────────────────────
        bullet = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline(p, bullet.group(2))
            indent = len(bullet.group(1))
            if indent:
                p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
            continue

        # ── Numbered lists ─────────────────────────────────────────────────
        numbered = re.match(r"^\s*\d+\.\s+(.*)", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            _parse_inline(p, numbered.group(1))
            continue

        # ── Normal paragraph ───────────────────────────────────────────────
        p = doc.add_paragraph(style="Normal")
        _parse_inline(p, line)


# ── Main builder ──────────────────────────────────────────────────────────────

def build_document() -> Path:
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Title page ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Super-Agent")
    title_run.bold = True
    title_run.font.size = Pt(32)
    _set_run_color(title_run, "1F3864")

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("Universal AI Agent — Project Documentation")
    sub_run.font.size = Pt(16)
    _set_run_color(sub_run, "2E4057")

    doc.add_paragraph()  # spacer

    desc_para = doc.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc_para.add_run(
        "A comprehensive reference covering system architecture, feature specifications,\n"
        "implementation plans, and the Auto-Evolution & Synthesized Consciousness Engine."
    )
    desc_run.font.size = Pt(11)
    _set_run_color(desc_run, "555555")

    _add_page_break(doc)

    # ── Table of Contents ────────────────────────────────────────────────────
    toc_heading = doc.add_heading("Table of Contents", level=1)
    _add_toc(doc)
    _add_page_break(doc)

    # ── Content sections ─────────────────────────────────────────────────────
    for section_title, filename in SECTIONS:
        md_path = DOCS_DIR / filename
        if not md_path.exists():
            continue

        # Section heading (Heading 1)
        doc.add_heading(section_title, level=1)

        markdown_text = md_path.read_text(encoding="utf-8")
        _add_markdown_section(doc, markdown_text)

        _add_page_break(doc)

    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    out = build_document()
    print(f"✅  Word document generated: {out}")
